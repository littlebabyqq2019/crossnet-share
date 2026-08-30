#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Extract text from Word documents for indexing
"""
import sys
import os


def extract_text_aspose(doc_path):
    """Extract text using Aspose.Words"""
    try:
        import aspose.words as aw
        
        # 加载文档
        doc = aw.Document(doc_path)
        
        # 转换为纯文本
        text = doc.to_string(aw.SaveFormat.TEXT)
        
        return text
    except Exception as e:
        raise Exception(f"Aspose.Words extraction failed: {e}")


def extract_text_python_docx(doc_path):
    """Extract text using python-docx (仅支持 .docx)"""
    try:
        from docx import Document
        
        doc = Document(doc_path)
        
        text_parts = []
        
        # 提取段落文本
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 提取表格文本
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        return '\n'.join(text_parts)
    except Exception as e:
        raise Exception(f"python-docx extraction failed: {e}")


def extract_text(doc_path):
    """
    Extract text from Word document
    Supports .doc and .docx formats
    """
    if not os.path.exists(doc_path):
        print(f"Error: File not found: {doc_path}", file=sys.stderr)
        return ""
    
    # 检查文件大小
    file_size_mb = os.path.getsize(doc_path) / (1024 * 1024)
    if file_size_mb > 50:
        print(f"Warning: Large Word file ({file_size_mb:.1f} MB), may take long time", file=sys.stderr)
    
    # 检查文件扩展名
    ext = os.path.splitext(doc_path)[1].lower()
    
    # 尝试多种提取方法
    methods = []
    
    # Aspose.Words 支持 .doc 和 .docx
    methods.append(("Aspose.Words", extract_text_aspose))
    
    # python-docx 只支持 .docx
    if ext == '.docx':
        methods.append(("python-docx", extract_text_python_docx))
    
    for method_name, method_func in methods:
        try:
            print(f"Trying {method_name}...", file=sys.stderr)
            text = method_func(doc_path)
            
            if text and len(text.strip()) > 10:
                print(f"Success with {method_name}, extracted {len(text)} characters", file=sys.stderr)
                return text
            else:
                print(f"{method_name} extracted no meaningful text", file=sys.stderr)
        except Exception as e:
            print(f"{method_name} failed: {e}", file=sys.stderr)
            continue
    
    # 所有方法都失败
    print("Error: All extraction methods failed", file=sys.stderr)
    print(f"Note: Make sure Aspose.Words is installed: pip install aspose-words", file=sys.stderr)
    return ""


def main():
    if len(sys.argv) < 2:
        print("Usage: extract_word_text.py <word_file>", file=sys.stderr)
        return 1
    
    doc_path = sys.argv[1]
    
    try:
        text = extract_text(doc_path)
        
        # 输出到stdout，让C++程序读取
        if text:
            print(text)
            return 0
        else:
            return 1
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        return 1


if __name__ == "__main__":
    sys.exit(main())
